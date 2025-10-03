import os
import time
import logging
import re
# Imports pour l'extraction de contenu
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from PIL import Image
    from PIL.ExifTags import TAGS
    IMAGE_AVAILABLE = True
except ImportError:
    IMAGE_AVAILABLE = False

# Configuration du logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('renommage_intelligent.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

def nettoyer_nom_fichier(nom_fichier):
    """
    Nettoie un nom de fichier en supprimant les caractères invalides.
    
    Args:
        nom_fichier: Nom original du fichier
        
    Returns:
        str: Nom nettoyé
    """
    # Caractères interdits dans les noms de fichiers
    caracteres_interdits = r'[<>:"/\\|?*]'
    nom_nettoye = re.sub(caracteres_interdits, '_', nom_fichier)
    
    # Supprimer les espaces multiples et les tirets multiples
    nom_nettoye = re.sub(r'\s+', ' ', nom_nettoye)
    nom_nettoye = re.sub(r'[_\-]+', '_', nom_nettoye)
    
    # Supprimer les espaces et underscores en début et fin
    nom_nettoye = nom_nettoye.strip(' _-')
    
    # Capitaliser les mots importants
    mots = nom_nettoye.split()
    mots_nettoyes = []
    for mot in mots:
        if len(mot) > 3 or mot.upper() in ['PDF', 'DOC', 'XLS']:
            mots_nettoyes.append(mot.capitalize())
        else:
            mots_nettoyes.append(mot.lower())
    
    return '_'.join(mots_nettoyes)

def extraire_contenu_pdf(chemin_fichier):
    """
    Extrait le contenu textuel d'un fichier PDF.
    
    Args:
        chemin_fichier: Chemin vers le fichier PDF
        
    Returns:
        str: Contenu textuel du PDF
    """
    contenu = ""
    
    if not PDF_AVAILABLE:
        logger.warning("PyPDF2 et pdfplumber non installés. Installation: pip install PyPDF2 pdfplumber")
        return contenu
    
    try:
        # Essayer avec pdfplumber d'abord (meilleur pour l'extraction de texte)
        with pdfplumber.open(chemin_fichier) as pdf:
            for page in pdf.pages[:3]:  # Limiter aux 3 premières pages
                texte_page = page.extract_text()
                if texte_page:
                    contenu += texte_page + "\n"
                if len(contenu) > 2000:  # Limiter la taille
                    break
    except Exception as e:
        logger.debug(f"Erreur pdfplumber pour {chemin_fichier}: {e}")
        
        # Fallback avec PyPDF2
        try:
            with open(chemin_fichier, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for i, page in enumerate(pdf_reader.pages[:3]):
                    contenu += page.extract_text() + "\n"
                    if len(contenu) > 2000:
                        break
        except Exception as e2:
            logger.warning(f"Impossible d'extraire le contenu PDF de {chemin_fichier}: {e2}")
    
    return contenu

def extraire_contenu_docx(chemin_fichier):
    """
    Extrait le contenu textuel d'un fichier Word (.docx).
    
    Args:
        chemin_fichier: Chemin vers le fichier Word
        
    Returns:
        str: Contenu textuel du document
    """
    contenu = ""
    
    if not DOCX_AVAILABLE:
        logger.warning("python-docx non installé. Installation: pip install python-docx")
        return contenu
    
    try:
        doc = DocxDocument(chemin_fichier)
        
        # Extraire le titre et les propriétés du document
        proprietes = doc.core_properties
        if proprietes.title:
            contenu += f"TITRE: {proprietes.title}\n"
        if proprietes.subject:
            contenu += f"SUJET: {proprietes.subject}\n"
        
        # Extraire le contenu des paragraphes
        for i, paragraph in enumerate(doc.paragraphs[:20]):  # Limiter aux 20 premiers paragraphes
            if paragraph.text.strip():
                contenu += paragraph.text + "\n"
                if len(contenu) > 2000:
                    break
        
        # Extraire le contenu des tableaux
        for table in doc.tables[:3]:  # Limiter aux 3 premiers tableaux
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        contenu += cell.text + " "
                contenu += "\n"
                if len(contenu) > 2000:
                    break
    
    except Exception as e:
        logger.warning(f"Impossible d'extraire le contenu Word de {chemin_fichier}: {e}")
    
    return contenu

def extraire_contenu_excel(chemin_fichier):
    """
    Extrait le contenu d'un fichier Excel.
    
    Args:
        chemin_fichier: Chemin vers le fichier Excel
        
    Returns:
        str: Contenu textuel du fichier Excel
    """
    contenu = ""
    
    if not EXCEL_AVAILABLE:
        logger.warning("openpyxl et pandas non installés. Installation: pip install openpyxl pandas")
        return contenu
    
    try:
        # Utiliser openpyxl pour les propriétés et pandas pour le contenu
        workbook = openpyxl.load_workbook(chemin_fichier, read_only=True)
        
        # Extraire les propriétés du document
        proprietes = workbook.properties
        if proprietes.title:
            contenu += f"TITRE: {proprietes.title}\n"
        if proprietes.description:
            contenu += f"DESCRIPTION: {proprietes.description}\n"
        
        # Extraire les noms des feuilles
        noms_feuilles = workbook.sheetnames
        contenu += f"FEUILLES: {', '.join(noms_feuilles)}\n"
        
        # Extraire quelques données de chaque feuille
        for nom_feuille in noms_feuilles[:3]:  # Limiter aux 3 premières feuilles
            try:
                df = pd.read_excel(chemin_fichier, sheet_name=nom_feuille, nrows=10)
                contenu += f"\nFEUILLE {nom_feuille}:\n"
                contenu += f"COLONNES: {', '.join(df.columns.astype(str))}\n"
                
                # Ajouter quelques valeurs
                for col in df.columns[:5]:  # Limiter aux 5 premières colonnes
                    valeurs_uniques = df[col].dropna().unique()[:3]  # 3 premières valeurs uniques
                    if len(valeurs_uniques) > 0:
                        contenu += f"{col}: {', '.join(str(v) for v in valeurs_uniques)}\n"
                
                if len(contenu) > 2000:
                    break
            except Exception as e:
                logger.debug(f"Erreur lecture feuille {nom_feuille}: {e}")
                continue
        
        workbook.close()
    
    except Exception as e:
        logger.warning(f"Impossible d'extraire le contenu Excel de {chemin_fichier}: {e}")
    
    return contenu

def extraire_metadonnees_image(chemin_fichier):
    """
    Extrait les métadonnées d'une image.
    
    Args:
        chemin_fichier: Chemin vers le fichier image
        
    Returns:
        str: Métadonnées de l'image
    """
    contenu = ""
    
    if not IMAGE_AVAILABLE:
        return contenu
    
    try:
        with Image.open(chemin_fichier) as image:
            # Informations de base
            contenu += f"FORMAT: {image.format}\n"
            contenu += f"TAILLE: {image.size[0]}x{image.size[1]}\n"
            contenu += f"MODE: {image.mode}\n"
            
            # Extraire les métadonnées EXIF
            exifdata = image.getexif()
            if exifdata:
                for tag_id in exifdata:
                    tag = TAGS.get(tag_id, tag_id)
                    data = exifdata.get(tag_id)
                    if isinstance(data, str) and len(data) < 100:
                        contenu += f"{tag}: {data}\n"
    
    except Exception as e:
        logger.debug(f"Impossible d'extraire les métadonnées de {chemin_fichier}: {e}")
    
    return contenu

def analyser_contenu_fichier(chemin_fichier):
    """
    Analyse le contenu d'un fichier selon son type.
    
    Args:
        chemin_fichier: Chemin complet vers le fichier
        
    Returns:
        str: Contenu analysé du fichier
    """
    _, extension = os.path.splitext(chemin_fichier)
    extension = extension.lower()
    
    contenu = ""
    
    if extension == '.pdf':
        contenu = extraire_contenu_pdf(chemin_fichier)
    elif extension in ['.docx', '.doc']:
        contenu = extraire_contenu_docx(chemin_fichier)
    elif extension in ['.xlsx', '.xls']:
        contenu = extraire_contenu_excel(chemin_fichier)
    elif extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
        contenu = extraire_metadonnees_image(chemin_fichier)
    else:
        # Pour les fichiers texte simples
        if extension in ['.txt', '.csv', '.json', '.xml']:
            try:
                with open(chemin_fichier, 'r', encoding='utf-8', errors='ignore') as f:
                    contenu = f.read(2000)  # Limiter à 2000 caractères
            except Exception as e:
                logger.debug(f"Impossible de lire {chemin_fichier}: {e}")
    
    return contenu

def extraire_mots_cles(contenu, langue='fr'):
    """
    Extrait les mots-clés pertinents du contenu d'un document.
    
    Args:
        contenu: Contenu textuel du document
        langue: Langue du document (fr/en)
        
    Returns:
        list: Liste des mots-clés les plus pertinents
    """
    if not contenu:
        return []
    
    # Mots vides en français et anglais
    mots_vides_fr = {
        'le', 'la', 'les', 'un', 'une', 'des', 'du', 'de', 'et', 'ou', 'mais', 'donc', 'car',
        'ce', 'cette', 'ces', 'dans', 'sur', 'avec', 'pour', 'par', 'sans', 'sous', 'vers',
        'chez', 'que', 'qui', 'quoi', 'dont', 'où', 'quand', 'comment', 'pourquoi',
        'il', 'elle', 'ils', 'elles', 'je', 'tu', 'nous', 'vous', 'mon', 'ma', 'mes',
        'ton', 'ta', 'tes', 'son', 'sa', 'ses', 'notre', 'votre', 'leur', 'leurs',
        'être', 'avoir', 'faire', 'dire', 'aller', 'voir', 'savoir', 'pouvoir',
        'vouloir', 'venir', 'falloir', 'devoir', 'prendre', 'donner', 'mettre'
    }
    
    mots_vides_en = {
        'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with',
        'by', 'from', 'up', 'about', 'into', 'through', 'during', 'before', 'after',
        'above', 'below', 'between', 'among', 'is', 'are', 'was', 'were', 'be', 'been',
        'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
        'should', 'may', 'might', 'must', 'can', 'this', 'that', 'these', 'those'
    }
    
    mots_vides = mots_vides_fr | mots_vides_en
    
    # Nettoyer et tokeniser le contenu
    contenu_nettoye = re.sub(r'[^\w\s]', ' ', contenu.lower())
    mots = contenu_nettoye.split()
    
    # Filtrer les mots
    mots_filtres = []
    for mot in mots:
        if (len(mot) >= 3 and 
            mot not in mots_vides and 
            not mot.isdigit() and
            not re.match(r'^[0-9]+$', mot)):
            mots_filtres.append(mot)
    
    # Compter les occurrences
    comptage_mots = {}
    for mot in mots_filtres:
        comptage_mots[mot] = comptage_mots.get(mot, 0) + 1
    
    # Trier par fréquence et prendre les plus fréquents
    mots_tries = sorted(comptage_mots.items(), key=lambda x: x[1], reverse=True)
    
    # Retourner les 5 mots les plus fréquents
    return [mot[0] for mot in mots_tries[:5]]

def generer_nom_intelligent(chemin_fichier):
    """
    Génère un nom de fichier intelligent basé sur le contenu.
    
    Args:
        chemin_fichier: Chemin complet vers le fichier
        
    Returns:
        str: Nouveau nom suggéré pour le fichier
    """
    nom_original = os.path.basename(chemin_fichier)
    nom_base, extension = os.path.splitext(nom_original)
    
    # Analyser le contenu
    contenu = analyser_contenu_fichier(chemin_fichier)
    
    if not contenu:
        logger.info(f"Aucun contenu extractible pour {nom_original}")
        return nom_original
    
    # Extraire les mots-clés
    mots_cles = extraire_mots_cles(contenu)
    
    if not mots_cles:
        logger.info(f"Aucun mot-clé trouvé pour {nom_original}")
        return nom_original
    
    # Chercher des patterns spécifiques dans le contenu
    patterns_specifiques = {
        'facture': r'facture|invoice|bill',
        'rapport': r'rapport|report|bilan',
        'contrat': r'contrat|contract|accord',
        'presentation': r'présentation|presentation|slide',
        'budget': r'budget|finance|cost',
        'planning': r'planning|schedule|agenda',
        'analyse': r'analyse|analysis|étude',
        'procedure': r'procédure|procedure|process',
        'manuel': r'manuel|manual|guide',
        'specification': r'spécification|specification|spec'
    }
    
    type_document = None
    for type_doc, pattern in patterns_specifiques.items():
        if re.search(pattern, contenu.lower()):
            type_document = type_doc
            break
    
    # Construire le nouveau nom
    elements_nom = []
    
    # Ajouter le type de document s'il est identifié
    if type_document:
        elements_nom.append(type_document.capitalize())
    
    # Ajouter les mots-clés les plus pertinents
    for mot_cle in mots_cles[:3]:  # Maximum 3 mots-clés
        if mot_cle not in (type_document or ''):
            elements_nom.append(mot_cle.capitalize())
    
    if not elements_nom:
        elements_nom = [mot_cle.capitalize() for mot_cle in mots_cles[:2]]
    
    # Construire le nom final
    if elements_nom:
        nouveau_nom_base = '_'.join(elements_nom)
        nouveau_nom_base = nettoyer_nom_fichier(nouveau_nom_base)
        
        # Limiter la longueur
        if len(nouveau_nom_base) > 50:
            nouveau_nom_base = nouveau_nom_base[:50]
        
        nouveau_nom = f"{nouveau_nom_base}{extension}"
    else:
        nouveau_nom = nom_original
    
    return nouveau_nom

def verifier_conflit_fichier(chemin_fichier):
    """
    Vérifie s'il y a un conflit de nom et génère un nom unique si nécessaire.
    
    Args:
        chemin_fichier: Chemin complet du fichier
        
    Returns:
        str: Chemin unique (modifié si nécessaire)
    """
    if not os.path.exists(chemin_fichier):
        return chemin_fichier
    
    # Séparer le répertoire, le nom et l'extension
    repertoire = os.path.dirname(chemin_fichier)
    nom_complet = os.path.basename(chemin_fichier)
    nom_base, extension = os.path.splitext(nom_complet)
    
    compteur = 1
    while os.path.exists(chemin_fichier):
        nouveau_nom = f"{nom_base}_({compteur}){extension}"
        chemin_fichier = os.path.join(repertoire, nouveau_nom)
        compteur += 1
        
        # Éviter une boucle infinie
        if compteur > 1000:
            logger.error(f"Trop de fichiers en conflit pour {nom_complet}")
            break
    
    return chemin_fichier

def renommer_fichiers(dossier, mode_simulation=False, limite_traitement=None, 
                     filtres_extension=None, exclure_motifs=None):
    """
    Renomme les fichiers selon leur contenu dans le dossier spécifié.
    
    Args:
        dossier (str): Le dossier contenant les fichiers à renommer
        mode_simulation (bool): Si True, montre les actions sans les exécuter
        limite_traitement (int): Nombre maximum de fichiers à traiter (None pour tous)
        filtres_extension (list): Liste des extensions à traiter (ex: ['.pdf', '.docx'])
        exclure_motifs (list): Liste de motifs à exclure du renommage
        
    Returns:
        dict: Résultats du traitement
    """
    # Vérifier que le dossier existe
    if not os.path.exists(dossier):
        logger.error(f"Le dossier {dossier} n'existe pas")
        return {'erreur': 'Dossier inexistant', 'fichiers_traites': 0}
    
    # Extensions supportées
    extensions_supportees = {'.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt', '.csv', 
                           '.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.json', '.xml'}
    
    # Obtenir la liste des fichiers
    tous_fichiers = [f for f in os.listdir(dossier) if os.path.isfile(os.path.join(dossier, f))]
    
    # Filtrer par extensions supportées si aucun filtre spécifique
    if not filtres_extension:
        fichiers_filtres = []
        for fichier in tous_fichiers:
            _, ext = os.path.splitext(fichier)
            if ext.lower() in extensions_supportees:
                fichiers_filtres.append(fichier)
        tous_fichiers = fichiers_filtres
        logger.info(f"Filtrage automatique: {len(tous_fichiers)} fichiers supportés trouvés")
    else:
        # Appliquer les filtres d'extension spécifiés
        fichiers_filtres = []
        for fichier in tous_fichiers:
            _, ext = os.path.splitext(fichier)
            if ext.lower() in [e.lower() for e in filtres_extension]:
                fichiers_filtres.append(fichier)
        tous_fichiers = fichiers_filtres
        logger.info(f"Filtrage par extension: {len(tous_fichiers)} fichiers sélectionnés")
    
    # Exclure les motifs si spécifiés
    if exclure_motifs:
        fichiers_filtres = []
        for fichier in tous_fichiers:
            exclure = False
            for motif in exclure_motifs:
                if motif.lower() in fichier.lower():
                    exclure = True
                    break
            if not exclure:
                fichiers_filtres.append(fichier)
        tous_fichiers = fichiers_filtres
        logger.info(f"Exclusion de motifs: {len(tous_fichiers)} fichiers restants")
    
    # Appliquer la limite si spécifiée
    if limite_traitement and len(tous_fichiers) > limite_traitement:
        logger.info(f"Limitation à {limite_traitement} fichiers sur {len(tous_fichiers)} au total")
        tous_fichiers = tous_fichiers[:limite_traitement]
    
    # Résultats du traitement
    resultats = {
        'fichiers_traites': 0,
        'fichiers_ignores': 0,
        'erreurs': 0,
        'renommages': [],
        'erreurs_details': []
    }
    
    logger.info(f"Début du {'simulation de ' if mode_simulation else ''}renommage intelligent...")
    
    for i, fichier in enumerate(tous_fichiers, 1):
        chemin_complet = os.path.join(dossier, fichier)
        
        try:
            # Générer le nouveau nom basé sur le contenu
            nouveau_nom = generer_nom_intelligent(chemin_complet)
            nouveau_chemin = os.path.join(dossier, nouveau_nom)
            
            # Vérifier s'il y a déjà un fichier avec ce nom
            nouveau_chemin = verifier_conflit_fichier(nouveau_chemin)
            nouveau_nom = os.path.basename(nouveau_chemin)
            
            # Vérifier si le renommage est nécessaire
            if fichier == nouveau_nom:
                logger.info(f"[{i}/{len(tous_fichiers)}] Pas besoin de renommer: {fichier}")
                resultats['fichiers_ignores'] += 1
                continue
            
            # Enregistrer l'action prévue
            action = {
                'ancien_nom': fichier,
                'nouveau_nom': nouveau_nom,
                'taille_fichier': os.path.getsize(chemin_complet)
            }
            
            if mode_simulation:
                logger.info(f"[{i}/{len(tous_fichiers)}] [SIMULATION] Renommage: {fichier} → {nouveau_nom}")
                resultats['renommages'].append(action)
                resultats['fichiers_traites'] += 1
            else:
                # Tenter le renommage réel
                try:
                    os.rename(chemin_complet, nouveau_chemin)
                    logger.info(f"[{i}/{len(tous_fichiers)}] Renommé: {fichier} → {nouveau_nom}")
                    resultats['renommages'].append(action)
                    resultats['fichiers_traites'] += 1
                    
                except Exception as e:
                    logger.warning(f"Erreur lors du renommage de {fichier}: {e}")
                    # Attendre et réessayer une fois
                    try:
                        time.sleep(1)
                        os.rename(chemin_complet, nouveau_chemin)
                        logger.info(f"[{i}/{len(tous_fichiers)}] Renommé après reprise: {fichier} → {nouveau_nom}")
                        resultats['renommages'].append(action)
                        resultats['fichiers_traites'] += 1
                    except Exception as e2:
                        erreur_msg = f"Échec définitif pour {fichier}: {e2}"
                        logger.error(erreur_msg)
                        resultats['erreurs'] += 1
                        resultats['erreurs_details'].append({
                            'fichier': fichier,
                            'erreur': str(e2)
                        })
        
        except Exception as e:
            erreur_msg = f"Erreur inattendue avec {fichier}: {e}"
            logger.error(erreur_msg)
            resultats['erreurs'] += 1
            resultats['erreurs_details'].append({
                'fichier': fichier,
                'erreur': str(e)
            })
    
    # Rapport final
    logger.info("="*60)
    logger.info("RAPPORT DE RENOMMAGE INTELLIGENT")
    logger.info("="*60)
    logger.info(f"Mode: {'SIMULATION' if mode_simulation else 'RÉEL'}")
    logger.info(f"Fichiers analysés: {len(tous_fichiers)}")
    logger.info(f"Fichiers traités: {resultats['fichiers_traites']}")
    logger.info(f"Fichiers ignorés: {resultats['fichiers_ignores']}")
    logger.info(f"Erreurs: {resultats['erreurs']}")
    
    if resultats['renommages']:
        logger.info("\nExemples de renommages:")
        for i, renommage in enumerate(resultats['renommages'][:5]):
            logger.info(f"  {i+1}. {renommage['ancien_nom']} → {renommage['nouveau_nom']}")
    
    if resultats['erreurs_details']:
        logger.info("\nDétail des erreurs:")
        for erreur in resultats['erreurs_details']:
            logger.info(f"  - {erreur['fichier']}: {erreur['erreur']}")
    
    return resultats

def installer_dependances():
    """
    Guide d'installation des dépendances nécessaires.
    """
    print("="*60)
    print("DÉPENDANCES NÉCESSAIRES POUR LE RENOMMAGE INTELLIGENT")
    print("="*60)
    print()
    print("Pour un fonctionnement optimal, installez les packages suivants:")
    print()
    print("# Pour les fichiers PDF:")
    print("pip install PyPDF2 pdfplumber")
    print()
    print("# Pour les fichiers Word (.docx):")
    print("pip install python-docx")
    print()
    print("# Pour les fichiers Excel:")
    print("pip install openpyxl pandas")
    print()
    print("# Pour les images:")
    print("pip install Pillow")
    print()
    print("# Installation complète:")
    print("pip install PyPDF2 pdfplumber python-docx openpyxl pandas Pillow")
    print()
    print("="*60)

# Exemple d'utilisation
if __name__ == "__main__":
    # Afficher le guide d'installation
    
    
    # Exemple d'utilisation
    dossier_test = input("\nEntrez le chemin du dossier à traiter (ou appuyez sur Entrée pour quitter): ").strip()
    
    if dossier_test and os.path.exists(dossier_test):
        print(f"\nAnalyse du dossier: {dossier_test}")
        
        # Mode simulation d'abord
        resultats_simulation = renommer_fichiers(
            dossier=dossier_test,
            mode_simulation=True,
            limite_traitement=10  # Limiter à 10 fichiers pour le test
        )
        
        if resultats_simulation['fichiers_traites'] > 0:
            confirmation = input(f"\nVoulez-vous procéder au renommage réel de {resultats_simulation['fichiers_traites']} fichiers? (oui/non): ").lower()
            
            if confirmation in ['oui', 'o', 'yes', 'y']:
                resultats_reel = renommer_fichiers(
                    dossier=dossier_test,
                    mode_simulation=False,
                    limite_traitement=10
                )
                print(f"\nRenommage terminé! {resultats_reel['fichiers_traites']} fichiers traités.")
            else:
                print("Renommage annulé.")
        else:
            print("Aucun fichier à traiter.")
    else:
        print("Dossier invalide ou inexistant.")